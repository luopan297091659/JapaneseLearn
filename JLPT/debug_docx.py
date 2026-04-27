#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document

doc = Document('test.docx')
print("=== Document structure ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and len(text) < 150:
        print(f"{i}: {text}")
    elif text:
        print(f"{i}: {text[:150]}...")
    if i > 50:
        break

print("\n=== Looking for section headers and question markers ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if any(x in text for x in ['第', '語彙', '読解', '聴解', '問題', '答案', '詳解']):
        print(f"{i}: {text[:100]}")
